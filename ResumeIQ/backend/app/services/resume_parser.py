import os
import re
import json
import zipfile
from typing import Optional
from pathlib import Path
from app.config import get_settings

settings = get_settings()


class ResumeParserService:

    @staticmethod
    def parse_pdf(file_path: str) -> dict:
        import pdfplumber
        issues = []
        text_pages = []
        page_limit = settings.max_pdf_pages
        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) > page_limit:
                raise ValueError(f"PDF has more than {page_limit} pages")
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                text_pages.append(page_text)
                tables = page.find_tables()
                if tables:
                    issues.append({"type": "table", "page": i + 1, "detail": f"Table detected at page {i+1}", "severity": "high"})
                textboxes = page.chars
                if textboxes:
                    x_coords = [c["x0"] for c in textboxes]
                    if x_coords:
                        unique_x = set(round(x, 0) for x in x_coords)
                        col_count = sum(1 for x in unique_x if x > 50)
                        if col_count > 1 and len(text_pages[i]) < 500:
                            issues.append({"type": "multi_column", "page": i + 1, "detail": f"Possible multi-column on page {i+1}", "severity": "high"})
                images = page.images
                if images:
                    issues.append({"type": "image", "page": i + 1, "detail": f"{len(images)} image(s) on page {i+1}", "severity": "medium"})
        raw_text = "\n".join(text_pages)
        max_chars = settings.max_paste_chars
        if len(raw_text) > max_chars:
            raise ValueError(f"Extracted text exceeds {max_chars} characters")
        ats_view = ResumeParserService._simulate_ats_view(text_pages)
        has_header_footer = ResumeParserService._detect_header_footer(text_pages)
        if has_header_footer:
            issues.append({"type": "header_footer", "detail": "Headers/footers detected", "severity": "medium"})
        font_issues = ResumeParserService._check_font_encoding(raw_text)
        issues.extend(font_issues)
        parsed = ResumeParserService._parse_sections(raw_text)
        return {"raw_text": raw_text, "ats_view_text": ats_view, "parsed_json": parsed, "has_parsing_issues": len(issues) > 0, "parsing_issues": json.dumps(issues)}

    @staticmethod
    def parse_docx(file_path: str) -> dict:
        max_uncompressed = settings.max_docx_uncompressed_mb * 1024 * 1024
        try:
            with zipfile.ZipFile(file_path) as zf:
                infos = zf.infolist()
                if len(infos) > 10_000:
                    raise ValueError("DOCX contains too many internal entries")
                total_uncompressed = sum(i.file_size for i in infos)
                if total_uncompressed > max_uncompressed:
                    raise ValueError(f"DOCX expands beyond {settings.max_docx_uncompressed_mb} MB when extracted")
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Invalid DOCX archive: {e}") from e
        from docx import Document
        issues = []
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if doc.tables:
            issues.append({"type": "table", "detail": f"{len(doc.tables)} table(s) detected", "severity": "high"})
        for para in doc.paragraphs:
            if para.style.name.startswith("Header"):
                issues.append({"type": "header_footer", "detail": "Headers detected", "severity": "medium"})
                break
        raw_text = "\n".join(paragraphs)
        max_chars = settings.max_paste_chars
        if len(raw_text) > max_chars:
            raise ValueError(f"Extracted text exceeds {max_chars} characters")
        ats_view = ResumeParserService._simulate_ats_view([raw_text])
        parsed = ResumeParserService._parse_sections(raw_text)
        return {"raw_text": raw_text, "ats_view_text": ats_view, "parsed_json": parsed, "has_parsing_issues": len(issues) > 0, "parsing_issues": json.dumps(issues)}

    @staticmethod
    def parse_text(text: str) -> dict:
        raw_text = text.strip()
        ats_view = ResumeParserService._simulate_ats_view([raw_text])
        parsed = ResumeParserService._parse_sections(raw_text)
        return {"raw_text": raw_text, "ats_view_text": ats_view, "parsed_json": parsed, "has_parsing_issues": False, "parsing_issues": "[]"}

    @staticmethod
    def _simulate_ats_view(text_pages: list) -> str:
        lines = []
        for page_text in text_pages:
            for line in page_text.split("\n"):
                stripped = line.strip()
                if stripped:
                    lines.append(stripped)
        return "\n".join(lines)

    @staticmethod
    def _detect_header_footer(text_pages: list) -> bool:
        patterns = [r"^\s*(page\s+\d+|p\.?\s*\d+)\s*$", r"^\s*\d+\s*$", r"^\s*resume\s*$", r"^\s*curriculum\s*vitae\s*$", r"^\s*cv\s*$", r"^[A-Z\s]{3,30}$"]
        for page_text in text_pages:
            lines = page_text.strip().split("\n")
            if len(lines) < 3:
                continue
            first, last = lines[0].strip(), lines[-1].strip()
            for pat in patterns:
                if re.match(pat, first, re.IGNORECASE) or re.match(pat, last, re.IGNORECASE):
                    return True
        return False

    @staticmethod
    def _check_font_encoding(text: str) -> list:
        issues = []
        special = re.findall(r'[^\x00-\x7F\u2022\u2023\u25E6\u25CB\u25A0\xB7\u2013\u2014\u2018\u2019\u201C\u201D\u2026]', text)
        if special:
            issues.append({"type": "special_characters", "detail": f"Non-standard chars: {''.join(set(special))[:10]}", "severity": "low"})
        return issues

    @staticmethod
    def _parse_sections(text: str) -> dict:
        section_keywords = {
            "contact": r"(?i)(contact|personal\s*info|contact\s*information)",
            "summary": r"(?i)(summary|professional\s*summary|profile|objective|about\s*me)",
            "skills": r"(?i)(skills|technical\s*skills|core\s*competencies|expertise|technologies)",
            "experience": r"(?i)(experience|work\s*experience|employment|professional\s*experience|career)",
            "education": r"(?i)(education|academic|qualifications|degrees?)",
            "certifications": r"(?i)(certifications|certificates|licenses|accreditations)",
            "projects": r"(?i)(projects|personal\s*projects|key\s*projects|portfolio)"
        }
        lines = text.split("\n")
        sections = {}
        current_section = "header"
        current_content = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                current_content.append("")
                continue
            matched = None
            for sname, pat in section_keywords.items():
                if re.match(pat, stripped) and len(stripped) < 60:
                    matched = sname
                    break
            if matched:
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = matched
                current_content = []
            else:
                current_content.append(stripped)
        if current_content:
            sections[current_section] = "\n".join(current_content)
        return ResumeParserService._structure_sections(sections)

    @staticmethod
    def _structure_sections(sections: dict) -> dict:
        contact_info = ResumeParserService._extract_contact_info(sections.get("header", ""))
        summary = sections.get("summary", "")
        skills_text = sections.get("skills", "")
        skills = []
        if skills_text:
            for sep in [",", "|", "\u2022", "\n"]:
                if sep in skills_text:
                    skills = [s.strip() for s in re.split(re.escape(sep), skills_text) if s.strip()]
                    break
            if not skills:
                skills = [skills_text]
        experience = ResumeParserService._parse_experience(sections.get("experience", ""))
        education = ResumeParserService._parse_education(sections.get("education", ""))
        certs_text = sections.get("certifications", "")
        certifications = [c.strip() for c in certs_text.split("\n") if c.strip()] if certs_text else []
        projects = ResumeParserService._parse_experience(sections.get("projects", ""))
        return {"contact_info": contact_info, "summary": summary, "skills": skills, "experience": experience, "education": education, "certifications": certifications, "projects": projects}

    @staticmethod
    def _extract_contact_info(header_text: str) -> dict:
        contact = {}
        email = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', header_text)
        if email:
            contact["email"] = email.group()
        phone = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', header_text)
        if phone:
            contact["phone"] = phone.group()
        linkedin = re.search(r'(linkedin\.com/in/[\w-]+)', header_text, re.IGNORECASE)
        if linkedin:
            contact["linkedin"] = linkedin.group()
        lines = header_text.split("\n")
        name = lines[0].strip() if lines else ""
        if name and len(name) < 60:
            contact["name"] = name
        loc = re.search(r'([A-Za-z\s]+,\s*[A-Z]{2})', header_text)
        if loc:
            contact["location"] = loc.group().strip()
        return contact

    @staticmethod
    def _parse_experience(text: str) -> list:
        if not text.strip():
            return []
        entries = []
        blocks = re.split(r'\n\s*\n', text)
        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue
            entry = {"title": lines[0], "company": "", "dates": "", "bullets": []}
            for line in lines[1:]:
                dm = re.match(r'([A-Za-z]{3,9}\s*\d{4})\s*[-u2013to]+\s*([A-Za-z]{3,9}\s*\d{4}|Present|Current|Now)', line, re.IGNORECASE)
                if dm:
                    entry["dates"] = line
                    continue
                if not entry["company"] and len(line) < 60:
                    entry["company"] = line
                    continue
                entry["bullets"].append(line.lstrip("\u2022-*0123456789). "))
            entries.append(entry)
        return entries

    @staticmethod
    def _parse_education(text: str) -> list:
        if not text.strip():
            return []
        entries = []
        blocks = re.split(r'\n\s*\n', text)
        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue
            entry = {"institution": lines[0], "degree": "", "dates": "", "details": []}
            for line in lines[1:]:
                dm = re.match(r'([A-Za-z]{3,9}\s*\d{4})\s*[-u2013to]+\s*([A-Za-z]{3,9}\s*\d{4}|Present|Expected)', line, re.IGNORECASE)
                if dm:
                    entry["dates"] = line
                    continue
                if not entry["degree"] and any(w in line.lower() for w in ["bachelor", "master", "phd", "b.s.", "m.s.", "b.a.", "m.a.", "associate", "diploma", "ph.d"]):
                    entry["degree"] = line
                else:
                    entry["details"].append(line)
            entries.append(entry)
        return entries

    @staticmethod
    def save_upload(file_bytes: bytes, filename: str) -> str:
        upload_dir = Path(settings.upload_dir)
        upload_dir.mkdir(parents=True, exist_ok=True)
        fp = upload_dir / filename
        with open(fp, "wb") as f:
            f.write(file_bytes)
        return str(fp)

    @staticmethod
    def parse_file(file_path: str) -> dict:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return ResumeParserService.parse_pdf(file_path)
        elif ext == ".docx":
            return ResumeParserService.parse_docx(file_path)
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            return ResumeParserService.parse_text(text)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
