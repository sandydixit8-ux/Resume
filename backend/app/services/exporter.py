"""Resume exporters: DOCX, PDF, HTML, Markdown, LaTeX, JSON Resume, Europass XML.

All exporters take the same parsed resume dict (from resume.parsed_json) plus
optional country rules, and return bytes + a filename.
"""

import io
import json
import xml.etree.ElementTree as ET
from app.services.country_rules import get_country

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except Exception:  # pragma: no cover
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    PDF_OK = True
except Exception:  # pragma: no cover
    PDF_OK = False


def _name(parsed):
    personal = parsed.get("personal", {}) or {}
    name = personal.get("name") or parsed.get("name") or ""
    if not name:
        first = (personal.get("first_name") or "").strip()
        last = (personal.get("last_name") or "").strip()
        name = f"{first} {last}".strip()
    return name or "Your Name"


def _contact_lines(parsed):
    personal = parsed.get("personal", {}) or {}
    phone = personal.get("phone") or parsed.get("phone") or ""
    email = personal.get("email") or parsed.get("email") or ""
    location = personal.get("location") or parsed.get("location") or ""
    links = [l for l in [personal.get("linkedin"), personal.get("github"), personal.get("website")] if l]
    return [x for x in [phone, email, location, *links] if x]


def _sections(parsed, country_code):
    """Return ordered [(title, [lines])] using country rules, filtering empties."""
    country = get_country(country_code)
    personal = parsed.get("personal", {}) or {}
    sections = []

    if country.get("summary") and parsed.get("summary"):
        sections.append(("Professional Summary", [parsed["summary"]]))

    skills = parsed.get("skills", [])
    if skills:
        sections.append(("Skills", [", ".join(skills)]))

    for exp in parsed.get("experience", []):
        title = exp.get("title", "").strip()
        company = exp.get("company", "").strip()
        dates = exp.get("dates", "").strip()
        header = title
        if company and dates:
            header = f"{title} — {company} ({dates})"
        elif company:
            header = f"{title} — {company}"
        elif dates:
            header = f"{title} ({dates})"
        bullets = [b for b in exp.get("bullets", []) if b.strip()]
        if header or bullets:
            sections.append((header, bullets))

    certs = parsed.get("certifications", []) or parsed.get("certificates", [])
    if certs:
        sections.append(("Certifications", [c if isinstance(c, str) else c.get("name", "") for c in certs]))

    edu = parsed.get("education", [])
    if edu:
        edu_lines = []
        for e in edu:
            degree = e.get("degree", "") or e.get("title", "")
            school = e.get("school", "") or e.get("institution", "")
            dates = e.get("dates", "")
            line = " ".join(x for x in [degree, school, dates] if x).strip()
            if line:
                edu_lines.append(line)
        if edu_lines:
            sections.append(("Education", edu_lines))

    langs = personal.get("languages") or parsed.get("languages") or []
    if langs:
        sections.append(("Languages", [l if isinstance(l, str) else f"{l.get('language')} ({l.get('level')})" for l in langs]))

    return sections


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def to_markdown(parsed: dict, country_code: str = "us") -> str:
    name = _name(parsed)
    lines = [f"# {name}"]
    contact = _contact_lines(parsed)
    if contact:
        lines.append(" | ".join(contact))
    for title, content in _sections(parsed, country_code):
        lines.append("")
        lines.append(f"## {title}")
        for item in content:
            lines.append(f"- {item}")
    return "\n".join(lines) + "\n"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def to_docx(parsed: dict, country_code: str = "us") -> bytes:
    if not DOCX_OK:
        raise RuntimeError("python-docx is not installed")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    name = _name(parsed)
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    contact = _contact_lines(parsed)
    if contact:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.add_run(" | ".join(contact))
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    for title, content in _sections(parsed, country_code):
        sh = doc.add_heading(title, level=2)
        for run in sh.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            run.font.size = Pt(12)
        for item in content:
            p = doc.add_paragraph(item, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

def to_html(parsed: dict, country_code: str = "us", template: str = "professional") -> str:
    name = _name(parsed)
    contact = _contact_lines(parsed)
    accent = {"modern": "#2563eb", "executive": "#b45309", "creative": "#db2777", "technology": "#0f172a"}.get(template, "#1f2937")
    body = []
    for title, content in _sections(parsed, country_code):
        body.append(f"<h2>{title}</h2>")
        body.append("<ul>" + "".join(f"<li>{item}</li>" for item in content) + "</ul>")
    return f"""<!doctype html><html lang="en"><head><meta charset="utf-8"><title>{name}</title>
<style>
body{{font-family:Georgia,serif;max-width:750px;margin:30px auto;color:#111;line-height:1.5}}
h1{{text-align:center;margin-bottom:2px}}
.contact{{text-align:center;color:#555;font-size:14px;margin-bottom:20px}}
h2{{color:{accent};border-bottom:1px solid #e5e7eb;padding-bottom:4px;font-size:16px;margin-top:22px}}
ul{{margin:6px 0 0 18px;padding:0}}li{{margin-bottom:4px}}
</style></head><body>
<h1>{name}</h1><div class="contact">{' | '.join(contact)}</div>
{''.join(body)}
</body></html>"""


# ---------------------------------------------------------------------------
# LaTeX
# ---------------------------------------------------------------------------

def to_latex(parsed: dict, country_code: str = "us") -> str:
    name = _name(parsed)
    contact = _contact_lines(parsed)
    parts = ["\\documentclass[10pt]{article}", "\\usepackage[utf8]{inputenc}", "\\usepackage[margin=2cm]{geometry}",
             "\\usepackage{enumitem}", "\\usepackage{hyperref}", "\\pagestyle{empty}", "\\begin{document}",
             f"\\begin{{center}}{{\\huge\\bfseries {_escape_latex(name)}}}\\\\[4pt]",
             f"{_escape_latex(' | '.join(contact))}\\end{{center}}"]
    for title, content in _sections(parsed, country_code):
        parts.append(f"\\section*{{{_escape_latex(title)}}}")
        items = "".join(f"\\item {_escape_latex(item)}\n" for item in content)
        parts.append(f"\\begin{{itemize}}[leftmargin=*,itemsep=1pt]{items}\\end{{itemize}}")
    parts.append("\\end{document}")
    return "\n".join(parts) + "\n"


def _escape_latex(text: str) -> str:
    return (text.replace("\\", r"\textbackslash{}").replace("&", r"\&").replace("%", r"\%")
                .replace("$", r"\$").replace("#", r"\#").replace("_", r"\_")
                .replace("{", r"\{").replace("}", r"\}").replace("~", r"\textasciitilde{}"))


# ---------------------------------------------------------------------------
# JSON Resume (jsonresume.org schema subset)
# ---------------------------------------------------------------------------

def to_json_resume(parsed: dict, country_code: str = "us") -> str:
    personal = parsed.get("personal", {}) or {}
    resume = {
        "basics": {
            "name": _name(parsed),
            "email": personal.get("email") or parsed.get("email", ""),
            "phone": personal.get("phone") or parsed.get("phone", ""),
            "location": {"city": personal.get("location") or parsed.get("location", "")},
            "summary": parsed.get("summary", ""),
            "profiles": [
                {"network": "linkedin", "url": personal.get("linkedin", "")},
                {"network": "github", "url": personal.get("github", "")},
            ],
        },
        "skills": [{"name": s} for s in parsed.get("skills", [])],
        "work": [
            {
                "name": e.get("company", ""),
                "position": e.get("title", ""),
                "startDate": (e.get("dates", "") or "").split("-")[0].strip(),
                "endDate": (e.get("dates", "") or "").split("-")[-1].strip() if "-" in (e.get("dates", "") or "") else "",
                "summary": "\n".join(e.get("bullets", [])),
            }
            for e in parsed.get("experience", [])
        ],
        "education": [
            {
                "institution": e.get("school") or e.get("institution", ""),
                "area": e.get("degree") or e.get("title", ""),
                "startDate": (e.get("dates", "") or "").split("-")[0].strip(),
            }
            for e in parsed.get("education", [])
        ],
    }
    return json.dumps(resume, indent=2)


# ---------------------------------------------------------------------------
# Europass XML (simplified)
# ---------------------------------------------------------------------------

def to_europass_xml(parsed: dict, country_code: str = "us") -> str:
    personal = parsed.get("personal", {}) or {}
    root = ET.Element("Europass", {
        "xmlns": "http://europass.cedefop.europa.eu/Europass",
        "version": "3.0",
    })
    profile = ET.SubElement(root, "SkillsPassport")
    learner = ET.SubElement(profile, "LearnerInfo")
    iden = ET.SubElement(learner, "Identification")
    person = ET.SubElement(iden, "PersonName")
    ET.SubElement(person, "FirstName").text = personal.get("first_name") or (_name(parsed).split(" ")[0])
    ET.SubElement(person, "Surname").text = personal.get("last_name") or " ".join(_name(parsed).split(" ")[1:]) or "-"
    contact = ET.SubElement(iden, "ContactInfo")
    ET.SubElement(contact, "Address").text = personal.get("location") or parsed.get("location", "")
    ET.SubElement(contact, "Email").text = personal.get("email") or parsed.get("email", "")
    ET.SubElement(contact, "Telephone").text = personal.get("phone") or parsed.get("phone", "")

    skills_section = ET.SubElement(learner, "Skills")
    for s in parsed.get("skills", []):
        ET.SubElement(skills_section, "Skill").text = s
    if parsed.get("summary"):
        ET.SubElement(learner, "Skills").text = ""

    work = ET.SubElement(learner, "WorkExperienceList")
    for e in parsed.get("experience", []):
        we = ET.SubElement(work, "WorkExperience")
        ET.SubElement(we, "Title").text = e.get("title", "")
        ET.SubElement(we, "Employer").text = e.get("company", "")
        ET.SubElement(we, "Period").text = e.get("dates", "")
        for b in e.get("bullets", []):
            ET.SubElement(we, "Description").text = b

    edu = ET.SubElement(learner, "EducationHistory")
    for e in parsed.get("education", []):
        ed = ET.SubElement(edu, "Education")
        ET.SubElement(ed, "Title").text = e.get("degree") or e.get("title", "")
        ET.SubElement(ed, "Organisation").text = e.get("school") or e.get("institution", "")
        ET.SubElement(ed, "Period").text = e.get("dates", "")

    return ET.tostring(root, encoding="unicode")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def to_pdf(parsed: dict, country_code: str = "us") -> bytes:
    if not PDF_OK:
        raise RuntimeError("reportlab is not installed")
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=18 * mm, rightMargin=18 * mm)
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=styles["Title"], fontName="Helvetica-Bold",
                                fontSize=20, alignment=1, textColor=colors.HexColor("#1f2937"))
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9.5,
                                   alignment=1, textColor=colors.HexColor("#4b5563"))
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold",
                              fontSize=12, textColor=colors.HexColor("#1f2937"),
                              spaceBefore=10, spaceAfter=2)
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10.5,
                                  leftIndent=12, spaceAfter=2)

    story = [Paragraph(_name(parsed), name_style)]
    contact = _contact_lines(parsed)
    if contact:
        story.append(Paragraph(" | ".join(contact), contact_style))
    for title, content in _sections(parsed, country_code):
        story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#e5e7eb"), spaceBefore=8, spaceAfter=2))
        story.append(Paragraph(title, h2_style))
        for item in content:
            story.append(Paragraph(item, bullet_style))
    doc.build(story)
    return buf.getvalue()


EXPORTERS = {
    "docx": (to_docx, ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf": (to_pdf, ".pdf", "application/pdf"),
    "html": (to_html, ".html", "text/html"),
    "md": (to_markdown, ".md", "text/markdown"),
    "markdown": (to_markdown, ".md", "text/markdown"),
    "tex": (to_latex, ".tex", "application/x-tex"),
    "latex": (to_latex, ".tex", "application/x-tex"),
    "json": (to_json_resume, ".json", "application/json"),
    "jsonresume": (to_json_resume, ".json", "application/json"),
    "europass": (to_europass_xml, ".xml", "application/xml"),
    "europass-xml": (to_europass_xml, ".xml", "application/xml"),
}


def export_resume(parsed: dict, fmt: str, country_code: str = "us", template: str = "professional"):
    """Return (bytes, filename, media_type). Raises ValueError for bad format."""
    key = (fmt or "").lower()
    if key not in EXPORTERS:
        raise ValueError(f"Unsupported format '{fmt}'")
    fn, ext, media = EXPORTERS[key]
    if key == "html":
        data = fn(parsed, country_code=country_code, template=template)
    else:
        data = fn(parsed, country_code=country_code)
    if isinstance(data, str):
        data = data.encode("utf-8")
    name = _name(parsed).replace(" ", "_") or "resume"
    filename = f"{name}{ext}"
    return data, filename, media
